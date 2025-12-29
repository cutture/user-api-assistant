from fastapi import FastAPI, HTTPException, Request
from pydantic import BaseModel
import uvicorn
import os
from dotenv import load_dotenv

# Load Env Vars FIRST
load_dotenv()

# Security & Rate Limiting
from slowapi import Limiter, _rate_limit_exceeded_handler
from slowapi.util import get_remote_address
from slowapi.errors import RateLimitExceeded
from core.security import configure_security
from utils.sanitization import sanitize_html

# Import the Graph (Now safe to import as env is loaded)
from agent.graph import app_graph
from typing import List, Dict, Any
from langchain_core.messages import HumanMessage, AIMessage

# Initialize Limiter
limiter = Limiter(key_func=get_remote_address)

app = FastAPI(
    title="Enterprise API Assistant",
    description="Agentic RAG for API Integration",
    version="1.0.0"
)

# Apply Rate Limit Exception Handler
app.state.limiter = limiter
app.add_exception_handler(RateLimitExceeded, _rate_limit_exceeded_handler)

# Apply Security Config (CORS, Headers)
# Apply Security Config (CORS, Headers)
configure_security(app)

# Register Advanced API Router
from api.endpoints import router as api_router
from api.sessions import router as session_router

app.include_router(api_router)
app.include_router(session_router)

class ChatRequest(BaseModel):
    query: str
    history: List[Dict[str, str]] = []

from core.exceptions import AppError, ServiceUnavailableError
from fastapi.responses import JSONResponse

@app.exception_handler(AppError)
async def app_exception_handler(request, exc: AppError):
    status_code = 500
    if isinstance(exc, ServiceUnavailableError):
        status_code = 503
    return JSONResponse(
        status_code=status_code,
        content={"error": exc.message, "details": exc.details, "type": exc.__class__.__name__},
    )

@app.get("/")
async def root():
    return {"status": "ok", "service": "Enterprise API Assistant", "version": "1.0.0"}

@app.get("/health")
@limiter.limit("5/minute") # Strict limit on health check to prevent spam
async def health_check(request: Request):
    # Perform deep check
    try:
        # Check Vector Store
        vector_store.client.heartbeat()
        return {"status": "ok", "service": "api-assistant-backend", "dependencies": {"chromadb": "ok"}}
    except Exception as e:
        return JSONResponse(status_code=503, content={"status": "degraded", "error": str(e)})

@app.post("/chat")
@limiter.limit("20/minute") # Initial limit for chat
async def chat(request: Request, body: ChatRequest): # Note: Body parameter must be explicit if Request is used
    """
    Invokes the Agentic Mesh (RAG -> Plan -> Code).
    """
    try:
        # Sanitization
        sanitized_query = sanitize_html(body.query)
        
        # Reconstruct Chat History
        messages = []
        for msg in body.history:
            if msg["role"] == "user":
                messages.append(HumanMessage(content=sanitize_html(msg["content"])))
            elif msg["role"] == "assistant":
                messages.append(AIMessage(content=sanitize_html(msg["content"])))
        
        # Add current query
        messages.append(HumanMessage(content=sanitized_query))

        inputs = {
            "messages": messages,
            "intent": "general",
            "context": [],
            "plan": "",
            "generated_code": "",
            "error": ""
        }
        
        # Invoke LangGraph
        result = app_graph.invoke(inputs)
        
        # Handle Reponse Logic
        response_content = result["generated_code"]
        plan_content = result["plan"]
        
        # ... logic continues ...
        if not response_content and "STATUS: INCOMPLETE" in plan_content:
             response_content = f"🛑 **Clarification Needed**\n\n{plan_content}"
        elif not response_content:
             response_content = "No code generated. Please check the Plan."

        return {
            "response": response_content,
            "plan": plan_content,
            "context": result["context"]
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

class PostmanRequest(BaseModel):
    plan: str
    code: str

@app.post("/export/postman")
async def export_postman(request: PostmanRequest):
    from utils.postman import convert_to_postman_collection
    try:
        return convert_to_postman_collection(request.plan, request.code)
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

# --- Document Upload Logic ---
from fastapi import UploadFile, File
import io
import pypdf
import uuid
from core.vector_store import store as vector_store
from core.text_splitter import APIDocSplitter

# Initialize services
# vector_store is imported as singleton
text_splitter = APIDocSplitter()

@app.post("/upload")
async def upload_document(file: UploadFile = File(...)):
    """
    Uploads a PDF or Text file, chunks it, and adds it to the Vector Store.
    """
    try:
        content_text = ""
        
        # 1. Extract Text
        if file.content_type == "application/pdf":
            pdf_bytes = await file.read()
            pdf_reader = pypdf.PdfReader(io.BytesIO(pdf_bytes))
            for page in pdf_reader.pages:
                text = page.extract_text()
                if text:
                    content_text += text + "\n"
        
        elif file.content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            import docx
            doc_bytes = await file.read()
            doc = docx.Document(io.BytesIO(doc_bytes))
            for para in doc.paragraphs:
                content_text += para.text + "\n"

        elif file.content_type == "application/json":
            content_bytes = await file.read()
            raw_text = content_bytes.decode("utf-8")
            # Create a nice string representation of the JSON for fallback
            import json
            json_data = json.loads(raw_text)
            content_text = json.dumps(json_data, indent=2)

        elif file.content_type in ["text/plain", "text/markdown", "text/csv"]:
            content_bytes = await file.read()
            content_text = content_bytes.decode("utf-8")

        else:
            raise HTTPException(status_code=400, detail=f"Unsupported file type: {file.content_type}. Use PDF, DOCX, JSON, or Text.")

        if not content_text.strip():
            raise HTTPException(status_code=400, detail="Empty file content.")

        # 3. Process Content using Parsers
        from core.parsers.factory import parser_factory
        from core.text_splitter import APIDocSplitter # Fallback
        
        # Determine content type for Parser Factory
        # We try to trust content-type, but fallback to extension if generic
        mime_type = file.content_type
        if mime_type == "application/octet-stream" or mime_type == "text/plain":
            if file.filename.endswith(".graphql") or file.filename.endswith(".gql"):
                mime_type = "application/graphql"
            elif file.filename.endswith(".json"):
                mime_type = "application/json"
        
        parser = parser_factory.get_parser(mime_type)
        
        documents_to_add = []
        metadatas_to_add = []
        
        if parser:
            print(f"🚀 Using Advanced Parser: {parser.__class__.__name__}")
            # Identify if content is json-like or text
            parse_input = content_text
            if mime_type == "application/json" or "json" in mime_type:
                 try:
                     import json
                     parse_input = json.loads(content_text)
                 except:
                     pass # Pass as string
            
            chunks = parser.parse(parse_input)
            documents_to_add = chunks
            metadatas_to_add = [{"source": file.filename, "type": parser.__class__.__name__} for _ in chunks]
            
        else:
            # Fallback to existing TextSplitter logic (PDF, Word, generic text)
            print("ℹ️ Using Standard Text Splitter")
            from langchain_core.documents import Document
            input_doc = Document(page_content=content_text, metadata={"source": file.filename})
            chunks = text_splitter.base_splitter.split_documents([input_doc])
            documents_to_add = [chunk.page_content for chunk in chunks]
            metadatas_to_add = [chunk.metadata for chunk in chunks]

        # 4. Add to Vector Store
        if documents_to_add:
            print(f"✅ Generated {len(documents_to_add)} chunks.")
            ids = [str(uuid.uuid4()) for _ in documents_to_add]
            vector_store.add_documents(documents=documents_to_add, metadatas=metadatas_to_add, ids=ids)

        return {
            "status": "success", 
            "filename": file.filename, 
            "chunks_added": len(documents_to_add),
            "total_chars": len(content_text)
        }

    except Exception as e:
        print(f"Upload Error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.delete("/reset")
async def reset_db():
    try:
        vector_store.reset()
        return {"status": "success", "message": "Vector Store cleared."}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

if __name__ == "__main__":
    uvicorn.run("main:app", host="0.0.0.0", port=8000, reload=True)
